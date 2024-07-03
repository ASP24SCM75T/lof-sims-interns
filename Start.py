import streamlit as st
import markdown2
import json
import requests
from sim_prompts import *  # Ensure this import provides the needed functionality
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document
from sqlalchemy import create_engine, Column, Integer, String, Text, MetaData, Index
from sqlalchemy.orm import sessionmaker, declarative_base

# Database setup
DATABASE_URL = "sqlite:///app_data.db"  # SQLite database

engine = create_engine(DATABASE_URL)
Session = sessionmaker(bind=engine)
session = Session()
Base = declarative_base()

# Define the models
class Transcript(Base):
    __tablename__ = 'transcripts'
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('transcript_content_idx', 'content'),)

class Assessment(Base):
    __tablename__ = 'assessments'
    id = Column(Integer, primary_key=True, autoincrement=True)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('assessment_content_idx', 'content'),)

class CaseDetails(Base):
    __tablename__ = 'case_details'
    id = Column(Integer, primary_key=True, autoincrement=True)
    saved_name = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    role = Column(String, nullable=False)
    specialty = Column(String, nullable=True)
    __table_args__ = (Index('case_details_content_idx', 'content'),)

# Create tables
Base.metadata.create_all(engine)

st.set_page_config(
    page_title='Simulated Case Generator',
    page_icon='🌌',
    layout="wide",
    initial_sidebar_state='auto'
)
with st.sidebar:
    st.image("static/er_bays.jpeg", use_column_width=True)

class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, self.title, 0, 1, "C")
        self.ln(10)

    def chapter_title(self, title, level=1):
        if level == 1:
            self.set_font("Arial", "B", 16)
            self.ln(10)
        elif level == 2:
            self.set_font("Arial", "B", 14)
            self.ln(8)
        elif level == 3:
            self.set_font("Arial", "B", 12)
            self.ln(6)
        self.cell(0, 10, title, 0, 1, "L")
        self.ln(2)

    def chapter_body(self, body):
        self.set_font("Arial", "", 12)
        self.multi_cell(0, 10, body)
        self.ln()

    def add_list(self, items, is_ordered=False):
        self.set_font("Arial", "", 12)
        for i, item in enumerate(items, start=1):
            if is_ordered:
                self.multi_cell(0, 10, f"{i}. {item}")
            else:
                self.multi_cell(0, 10, f"- {item}")
        self.ln()

def html_to_pdf(html_content, name):
    # Use BeautifulSoup to parse the HTML
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Extract title for the document
    case_title_tag = soup.find("h1")
    case_title = case_title_tag.get_text() if case_title_tag else "Document"
    
    # Create PDF instance with dynamic title
    pdf = PDF()
    pdf.title = case_title
    pdf.add_page()

    # Process each section of the HTML
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h1":
            pdf.chapter_title(element.get_text(), level=1)
        elif element.name == "h2":
            if "Patient Door Chart" in element.get_text():
                pdf.add_page()
            pdf.chapter_title(element.get_text(), level=2)
        elif element.name == "h3":
            pdf.chapter_title(element.get_text(), level=3)
        elif element.name == "p":
            pdf.chapter_body(element.get_text())
        elif element.name == "ul":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=False)
        elif element.name == "ol":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=True)
        elif element.name == "hr":
            pdf.add_page()
    
    # Output the PDF
    pdf.output(name)

def html_to_docx(html_content, name):
    # Use BeautifulSoup to parse the HTML
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Create DOCX instance
    doc = Document()

    # Extract title for the document
    case_title_tag = soup.find("h1")
    case_title = case_title_tag.get_text() if case_title_tag else "Document"
    doc.add_heading(case_title, level=1)
    
    # Process each section of the HTML
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(f"- {li.get_text()}")
        elif element.name == "ol":
            for i, li in enumerate(element.find_all("li"), start=1):
                doc.add_paragraph(f"{i}. {li.get_text()}")
        elif element.name == "hr":
            doc.add_page_break()
    
    # Output the DOCX
    doc.save(name)

def init_session():
    if "final_case" not in st.session_state:
        st.session_state["final_case"] = ""
    if "retrieved_case" not in st.session_state:
        st.session_state["retrieved_case"] = ""
    if "retrieved_name" not in st.session_state:
        st.session_state["retrieved_name"] = ""
    if "selected_case_id" not in st.session_state:
        st.session_state["selected_case_id"] = -1
    if "lab_tests" not in st.session_state:
        st.session_state["lab_tests"] = []
    if "selected_tests" not in st.session_state:
        st.session_state["selected_tests"] = []

# Function to save a transcript
def save_transcript(transcript_content, role, specialty):
    new_transcript = Transcript(content=transcript_content, role=role, specialty=specialty)
    session.add(new_transcript)
    session.commit()

# Function to save an assessment
def save_assessment(assessment_content, role, specialty):
    new_assessment = Assessment(content=assessment_content, role=role, specialty=specialty)
    session.add(new_assessment)
    session.commit()

# Function to save case details
def save_case_details(case_details_content, saved_name, role = "", specialty=""):
    new_case_details = CaseDetails(content=case_details_content, saved_name=saved_name, role=role, specialty=specialty)
    session.add(new_case_details)
    session.commit()

# Function to retrieve records with full-text search and wildcards
def get_records(model, search_text=None, saved_name=None, role=None, specialty=None):
    query = session.query(model)
    if search_text:
        search_text = f"%{search_text}%"  # Wildcard search
        query = query.filter(model.content.ilike(search_text))
    if saved_name:
        saved_name = f"%{saved_name}%"  # Wildcard search
        query = query.filter(model.saved_name.ilike(saved_name))
    if role:
        query = query.filter_by(role=role)
    if specialty:
        query = query.filter_by(specialty=specialty)
    return query.all()

def llm_call(model, messages):
    
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": "Bearer " + st.secrets["OPENROUTER_API_KEY"],  # Fixed to correct access to secrets
                "Content-Type": "application/json",
                "HTTP-Referer": "https://fsm-gpt-med-ed.streamlit.app",  # To identify your app
                "X-Title": "lof-sims",
            },
            data=json.dumps({
                "model": model,
                "messages": messages,
            })
        )
    except requests.exceptions.RequestException as e:
        st.error(f"Error - make sure bills are paid!: {e}")
        return None
    # Extract the response content
    try:
        response_data = response.json()
    except json.JSONDecodeError:
        st.error(f"Error decoding JSON response: {response.content}")
        return None
    return response_data  # Adjusted to match expected JSON structure

def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.write("*Please contact David Liebovitz, MD if you need an updated password for access.*")
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct.
        return True

st.title("Case Generator for Simulations")
init_session()

lab_tests_dict = {
    "Thyroid": {
        "TSH": "0.3-3.7 uIU/ml",
        "Free T3": "230-619 pg/dL",
        "Free T4": "0.7-1.9 ng/dL"
    },
    "Cardiac": {
        "CPK": "26-140 U/L",
        "CK-MB": "< 3% of total",
        "Troponin I": "0.00-0.04 ng/ml"
    },
    "CMP (Comprehensive Metabolic Panel)": {
        "Sodium": "135-145 mmol/L",
        "Potassium": "3.5-4.5 mmol/L",
        "Chloride": "98 - 106 mmol/L",
        "CO2": "22-29 mmol/L",
        "BUN": "7 - 18 mg/dL",
        "Cr": "0.6-1.2 mg/dL",
        "Glucose": "20-115 mg/dL",
        "Calcium": "8.4-10.2 mg/dL",
        "Magnesium": "1.3-2.1 mmol/L",
        "Alk Phos": "38-126 U/L",
        "Albumin": "3.5-5.5 g/dL",
        "Total Protein": "6 - 8 g/dL",
        "AST": "5 - 30 U/L",
        "ALT": "5 - 35 U/L",
        "Bilirubin": "< 1.0 mg/dL"
    }
    # Add other diagnoses and their tests here
}

if check_password():
    st.info("Provide inputs and generate a case. After your case is generated, please click the '*Send case to the simulator!*' and then wake the simulator.")
    
    with st.expander("Model Options for Case Generation (Claude3 Haiku by default)", expanded=False):
        model_choice = st.selectbox("Model Options", (
            "anthropic/claude-3-haiku",
            "anthropic/claude-3-sonnet", 
            "anthropic/claude-3-opus", 
            "openai/gpt-4o", 
            "google/gemini-pro", 
            "meta-llama/llama-2-70b-chat",
        ), index=0)

    if "response_markdown" not in st.session_state:
        st.session_state["response_markdown"] = ""
        
    if "expanded" not in st.session_state:
        st.session_state["expanded"] = True
        
    if "edited_new_case" not in st.session_state:
        st.session_state["edited_new_case"] = ""
        
    if "learner_tasks" not in st.session_state:
        st.session_state["learner_tasks"] = learner_tasks
        
    tab1, tab2, tab3 = st.tabs(["New Case", "Retrieve a Case", "Select Lab Tests"])
    
    with tab1:

        col1, col2, col3 = st.columns([2, 2, 5])

        with col1:    
            st.info("**Include desired history in the text paragraph. The AI will generate additional details as needed to draft an educational case.**")
                
            case_study_input = {
                'Case Title': st.text_input("Case Study Title", help="Presenting symptom, e.g."),
                'Case Description': st.text_area("Case Description", height=200, help = "As detailed or brief as desired, e.g., 65F with acute chest pain..."),
                'Case Primary Diagnosis': st.text_input("Primary Diagnosis", help = "The one or more primary diagnoses, e.g., Pulmonary Embolism"),
            }
            case_study_input = json.dumps(case_study_input)
            with st.expander("Default Learner Tasks", expanded = False):
                st.markdown(learner_tasks)  # Display the default tasks
            if st.checkbox("Edit Learner Tasks", value=False, key = "initial_case_edit"):
                learner_tasks = st.text_area("Learner Tasks for Assessment", height=200, help = "What the learner is expected to do, e.g., Perform a focused history and examination", value = learner_tasks)
            st.session_state.learner_tasks = learner_tasks
        
        with col1: 
            st.info("Click submit when ready to generate a case!")
            submit_button = st.button("Submit")

        if submit_button:
            messages = [
                {"role": "system", "content": f'Using the case details provided (supplemented as needed with additional generated content), comprehensively populate an educational clinical case description in the following specific format: {output_format}'},
                {"role": "user", "content": f'case_details: {case_study_input}'}
            ]
        
            with col2:
                with st.spinner("Assembling Case... Please wait."):
                    response_content = llm_call(model_choice, messages)
            st.session_state.response_markdown = response_content['choices'][0]['message']['content']
        if st.session_state.response_markdown != "":
            with col3:
                st.success("Case Study Framework Generated!")
                with st.expander("View Full Case", expanded=st.session_state.expanded):
                    st.markdown(st.session_state.response_markdown)
                    st.session_state.expanded = False
            
            with col2:
                st.info("Review and/or edit the case and begin the simulator!")                  
            
                if st.checkbox("Edit Case (Scroll Down)", value=False):
                    with col3:
                        st.session_state.expanded = False
                        st.warning("Please edit the case as needed while leaving other characters, e.g., '#' and '*', in place. Use the 'Save Case Edits' button at the bottom to save edits!")
                        
                        edited_new_case = st.text_area("Click button at bottom to save your edits!", st.session_state.response_markdown, height=1000) 
                        if st.button("Save Case Edits for the Simulator"):
                            st.success("Case Edits Saved!")
                            if edited_new_case:
                                st.session_state["final_case"] = edited_new_case
                        st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator (including any saved edits)", icon="🧠")
                else:
                    st.session_state["final_case"] = st.session_state.response_markdown
                
                if st.session_state.final_case !="":        
                    case_html = markdown2.markdown(st.session_state.final_case, extras=["tables"])
                        
                    if st.checkbox("Generate Case PDF file"):
                        html_to_pdf(case_html, 'case.pdf')
                        with open("case.pdf", "rb") as f:
                            st.download_button("Download Case PDF", f, "case.pdf")
                    if st.checkbox("Generate Case DOCX file"):
                        html_to_docx(case_html, 'case.docx')
                        with open("case.docx", "rb") as f:
                            st.download_button("Download Case DOCX", f, "case.docx")

                if st.session_state["final_case"] != "":
                    if st.button("Send case to the simulator!"):
                        st.session_state["final_case"] = st.session_state.final_case
                        st.session_state["retrieved_name"] = st.session_state.retrieved_name
                        st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator", icon="🧠")

        with col3:
            roles = ["1st year medical student", "2nd year medical student", "3rd year medical student", "4th year medical student", "Resident", "Fellow", "Attending"]
            if st.session_state.final_case:
                st.divider()
                if st.checkbox("Save Case to the Database for Future Use"):
                    case_details = st.text_area("Case Details to Save to the Database for Future Use", value=st.session_state.final_case)
                    saved_name = st.text_input("Saved Name (Required to save case)")

                    if st.button("Save Case to the Database for future use!"):
                        if saved_name:
                            save_case_details(case_details, saved_name)
                            st.success("Case Details saved successfully!")
                        else:
                            st.error("Saved Name is required to save the case")
    
    if "search_results" not in st.session_state:
        st.session_state.search_results = []
    if "selected_case" not in st.session_state:
        st.session_state.selected_case = None

    # Tab2 content for retrieving and selecting cases
    with tab2:
        
        col3, col4 = st.columns([1,3])
        with col3:
            st.header("Retrieve Records")
            search_text = st.text_input("Search Text")
            search_saved_name = st.text_input("Search by Saved Name")
            search_role =""
            search_specialty = ""
            if search_role in ["Resident", "Fellow", "Attending"]:
                search_specialty = st.text_input("Search by Specialty", "")

            if st.button("Search Cases"):
                st.session_state.search_results = get_records(CaseDetails, search_text, search_saved_name)

            if st.session_state.search_results:
                st.subheader("Cases Found")
                for i, case in enumerate(st.session_state.search_results):
                    st.write(f"{i+1}. {case.saved_name}")
                    if st.button(f"View (and Select) Case {i+1}", key=f"select_case_{i}"):
                        st.session_state.selected_case = case
            with col4:
                if st.session_state.selected_case:
                    st.subheader("Retrieved Case")
                    with st.expander("View Full Case", expanded=False):
                        st.write(f'Here is the retrieved case name: {st.session_state.selected_case.saved_name}')
                        st.write(st.session_state.selected_case.content)
                        st.session_state.final_case = st.session_state.selected_case.content
                    if st.checkbox("Edit Retrieved Case (Scroll Down)", value=False, key = "retrieved_case_edit"):
                        st.session_state.expanded = False
                        st.warning('Please edit the case as needed while leaving other characters, e.g., "#" and "*", in place. Remember to update the Door Chart section at the bottom!')
                        updated_retrieved_case = st.text_area("Edit Case, enter control-enter or command-enter to save edits!", st.session_state.selected_case.content, height=1000)
                        make_new_entry = st.checkbox("If desired, make a database entry when saving edits.", value=False)
                        if make_new_entry:
                            saved_name = st.text_input("Saved Name after edits (Required to save case)")
                        if st.button("Save Edits for Simulation and Generate a PDF"):
                            st.session_state.final_case = updated_retrieved_case
                            st.info("Case Edits Saved!")   
                            updated_case_html = markdown2.markdown(st.session_state.final_case, extras=["tables"])
                            html_to_pdf(updated_case_html, 'updated_case.pdf')
                            with open("updated_case.pdf", "rb") as f:
                                st.download_button("Download Updated Case PDF", f, "updated_case.pdf")
                            if st.button("Generate Case DOCX file"):
                                html_to_docx(updated_case_html, 'updated_case.docx')
                                with open("updated_case.docx", "rb") as f:
                                    st.download_button("Download Case DOCX", f, "updated_case.docx")
                            if make_new_entry:
                                if saved_name:
                                    save_case_details(st.session_state.final_case, saved_name)
                                    st.success("Case Details saved successfully!")
                                else:
                                    st.error("Saved Name is required to save the case")
                    st.page_link("pages/🧠_Simulator.py", label="Wake the Simulator ", icon="🧠")

    # Tab3 content for selecting lab tests
    with tab3:
        st.header("Select Lab Tests")
        primary_diagnosis = st.selectbox("Select Primary Diagnosis for Lab Tests", options=list(lab_tests_dict.keys()))
        
        if primary_diagnosis:
            st.session_state.lab_tests = lab_tests_dict.get(primary_diagnosis, {})

        if st.session_state.lab_tests:
            selected_tests = st.multiselect("Select Lab Tests", options=list(st.session_state.lab_tests.keys()))
            st.session_state.selected_tests = selected_tests

            if selected_tests:
                st.subheader("Selected Lab Tests")
                table_data = [["Test", "Result", "Normal Range", "Comments"]]
                for test in selected_tests:
                    table_data.append([test, "", st.session_state.lab_tests[test], ""])

                st.write("## Lab Test Results")
                st.table(table_data)
