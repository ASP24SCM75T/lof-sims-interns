import streamlit as st
from sim_prompts import *
import markdown2
from groq import Groq
from openai import OpenAI
import os
from bs4 import BeautifulSoup
from fpdf import FPDF
from datetime import datetime
from audio_recorder_streamlit import audio_recorder
from prompts import *
import tempfile
import requests
import json
import base64
import random
from Start import llm_call, PDF 
from docx import Document
from fpdf import FPDF

CHECKLIST_FIELDS = {
    "Onset": "",
    "Location": "",
    "Duration": "",
    "Character": "",
    "Aggravating/Alleviating factors": "",
    "Radiation": "",
    "Timing": "",
    "Severity": ""
}

QUESTION_TO_FIELD_MAPPING = {
    "onset": "Onset",
    "location": "Location",
    "duration": "Duration",
    "character": "Character",
    "aggravating": "Aggravating/Alleviating factors",
    "alleviating": "Aggravating/Alleviating factors",
    "radiation": "Radiation",
    "timing": "Timing",
    "severity": "Severity"
}

def assign_random_voice(sex):
    male_voices = ['echo', 'fable', 'onyx']
    female_voices = ['nova', 'shimmer']
    
    if sex == 'male':
        voices = male_voices
    else:
        voices = female_voices
    
    voice = random.choice(voices)
    return voice

def transcript_to_pdf(html_content, name):   
    html_content = html_content.replace('🤒', 'Patient').replace('👩‍⚕️', 'Doctor')
    html_content = html_content.encode('latin-1', 'ignore').decode('latin-1')
    soup = BeautifulSoup(html_content, "html.parser")
    title = "Patient Case"
    pdf = PDF()
    pdf.title = title
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    for element in soup.find_all(["h2", "h3", "p", "ul", "ol", "li", "hr"]):
        if element.name == "h2":
            pdf.chapter_title(element.get_text(), level=2)
        elif element.name == "h3":
            pdf.chapter_title(element.get_text(), level=3)
        elif element.name == "p":
            pdf.chapter_body(element.get_text())
        elif element.name == "ul":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=False)
        elif element.name == "ol":
            items = [li.get_text() for li in element.find_all("li")]
            pdf.add_list(items, is_ordered=True)
        elif element.name == "hr":
            pdf.add_page()
    
    pdf.output(name, 'F')


def html_to_pdf(html_content, name):
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        title = "Checklist"
        pdf = PDF()
        pdf.title = title
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for element in soup.find_all(["h2", "h3", "p", "ul", "ol", "li", "hr"]):
            if element.name == "h2":
                pdf.chapter_title(element.get_text(), level=2)
            elif element.name == "h3":
                pdf.chapter_title(element.get_text(), level=3)
            elif element.name == "p":
                pdf.chapter_body(element.get_text())
            elif element.name == "ul":
                items = [li.get_text() for li in element.find_all("li")]
                pdf.add_list(items, is_ordered=False)
            elif element.name == "ol":
                items = [li.get_text() for li in element.find_all("li")]
                pdf.add_list(items, is_ordered=True)
            elif element.name == "hr":
                pdf.add_page()
        
        pdf.output(name, 'F')
        return name
    except Exception as e:
        st.error(f"Failed to create PDF: {e}")
        return None
#sprint 3
def generate_combined_doc(checklist_html, assessment_html, orders_html, filename):
    doc = Document()

    def clean_html(html_content):
        soup = BeautifulSoup(html_content, "html.parser")
        return soup.get_text(separator="\n")
    def add_section_heading(heading):
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, heading, ln=True, align='C')
        pdf.set_font("Arial", size=12)
        pdf.ln(5)  

    # Add checklist content
    if checklist_html:
        add_section_heading("Checklist")
        cleaned_checklist = clean_html(checklist_html)
        doc.add_paragraph(cleaned_checklist)

    # Add assessment content
    if assessment_html:
        add_section_heading("Assessment")
        cleaned_assessment = clean_html(assessment_html)
        doc.add_paragraph(cleaned_assessment)

    # Add orders content
    if orders_html:
        add_section_heading("Orders")
        cleaned_orders = clean_html(orders_html)
        doc.add_paragraph(cleaned_orders)

    doc.save(filename)
    return filename
#sprint 3
def generate_combined_pdf(checklist_html, assessment_html, orders_html, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    def clean_html(html_content):
        soup = BeautifulSoup(html_content, "html.parser")
        return soup.get_text(separator="\n")

    def add_section_heading(heading):
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, heading, ln=True, align='C')
        pdf.set_font("Arial", size=12)
        pdf.ln(5)  

    if checklist_html:
        add_section_heading("Checklist")
        cleaned_checklist = clean_html(checklist_html)
        pdf.multi_cell(0, 4, cleaned_checklist)
    if assessment_html:
        pdf.add_page()
        add_section_heading("Assessment")
        cleaned_assessment = clean_html(assessment_html)
        pdf.multi_cell(0, 4, cleaned_assessment)
    if orders_html:
        pdf.add_page()
        add_section_heading("Orders")
        cleaned_orders = clean_html(orders_html)
        pdf.multi_cell(0, 4, cleaned_orders)

    pdf.output(filename)
    return filename

def transcribe_audio(audio_file_path):
    from openai import OpenAI
    api_key = st.secrets["OPENAI_API_KEY"]
    client = OpenAI(base_url="https://api.openai.com/v1", api_key=api_key)
    audio_file = open(audio_file_path, "rb")
    transcript = client.audio.transcriptions.create(
        model="whisper-1", 
        file=audio_file, 
        response_format="text"
    )
    return transcript

def talk_stream(model, voice, input):
    api_key = st.secrets["OPENAI_API_KEY"]
    client = OpenAI(base_url="https://api.openai.com/v1", api_key=api_key)
    try:
        response = client.audio.speech.create(
            model= model,
            voice= voice,
            input= input,
        )
        response.stream_to_file("last_interviewer.mp3")
    except Exception as e:
        st.write("The API is busy - should work in a moment for voice.")
    
def autoplay_local_audio(filepath: str):
    with open(filepath, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    md = f"""
        <audio controls autoplay="true">
        <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
        </audio>
        """
    st.markdown(md, unsafe_allow_html=True)

def parse_groq_stream(stream):
    for chunk in stream:
        if chunk.choices:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content

@st.cache_data
def extract_patient_door_chart_section(text):
    start_marker = "## PATIENT DOOR CHART"
    start_index = text.find(start_marker)
    if start_index != -1:
        return text[start_index:]
    else:
        return "PATIENT DOOR CHART section not found in the provided text. Please go back to the Start page!"

def infer_checklist_field(question):
    prompt = f"Which field does the following question relate to?\n\nQuestion: {question}\n\nFields: {', '.join(CHECKLIST_FIELDS.keys())}\n\nAnswer with the most relevant field:, the response should only contain the relevant feild name"
    response = llm_call("anthropic/claude-3-haiku", [{"role": "user", "content": prompt}])
    inferred_field = response['choices'][0]['message']['content'].strip()
    print(inferred_field)
    if inferred_field in CHECKLIST_FIELDS:
        return inferred_field
    return None

def update_checklist_with_answer(question, answer, checklist_fields):
    field = infer_checklist_field(question)
    if field:
        checklist_fields[field] = answer
    else:
        print("No Match detected")
 

def generate_checklist_template(checklist_fields):
    template = "HPI:\n\n"
    for field, value in checklist_fields.items():
        template += f"{field}: {value}\n\n"
    return template

try:
    extracted_section = extract_patient_door_chart_section(st.session_state.final_case)
    st.info(extracted_section)
    st.info(st.session_state.learner_tasks)
    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "system", "content": f'{sim_persona} Here are the specifics for your persona: {st.session_state.final_case}'}, ]
except Exception as e:
    st.error(f"Please return to the main page. An error occurred. Please do not 're-load' when in the middle of a conversation. Here are the error details: {e}. ")

if "sex" not in st.session_state:
    st.session_state.sex = ""
if st.session_state.sex == "":
    messages_sex = [{"role": "user", "content": f'Analyze the following content and return only the sex, e.g., male, female, or other. Return nothing else. {extracted_section}'}]
    st.session_state.sex = llm_call("anthropic/claude-3-haiku", messages_sex)

groq_client = Groq(api_key = st.secrets['GROQ_API_KEY'])

st.title("Clinical Simulator Chat")

if "password_correct" not in st.session_state:
    st.session_state["password_correct"] = False
    st.write("Login on the Sims page to get started.")

if st.session_state["password_correct"] == True:
    st.info("Type your questions at the bottom of the page or use voice input (left sidebar)! You may need to right click your Chrome browser tab to unmute this website and also accept the microphone permissions.")
    
    with st.sidebar:
        with st.expander("Change Model", expanded=False):
            st.session_state.model = st.selectbox(
                'voice a model',
                ['llama3-70b-8192', 'gpt-4o',], index=1,
            )

        if "checklist_fields" not in st.session_state:
            st.session_state.checklist_fields = CHECKLIST_FIELDS.copy()
            st.session_state.checklist_template = generate_checklist_template(st.session_state.checklist_fields)

        st.sidebar.subheader("Checklist")
        st.sidebar.markdown(st.session_state.checklist_template)

        if st.sidebar.button("Update and generate PDF"):
            checklist_html = markdown2.markdown(st.session_state.checklist_template, extras=["tables"])
            pdf_path = html_to_pdf(checklist_html, 'checklist.pdf')
            if pdf_path:
                with open(pdf_path, "rb") as f:
                    st.download_button("Download Checklist PDF", f, "checklist.pdf")
            else:
                st.error("Failed to create checklist PDF.")

    if "sim_response" not in st.session_state:
        st.session_state["sim_response"] = ""

    if "audio_off" not in st.session_state:
        st.session_state["audio_off"] = False

    if "audio_input" not in st.session_state:
        st.session_state["audio_input"] = ""
        
    if "voice" not in st.session_state:
        st.session_state["voice"] = assign_random_voice(st.session_state.sex)
        
    if "results" not in st.session_state:
        st.session_state["results"] = ""
        
    if "orders_placed" not in st.session_state:
        st.session_state["orders_placed"] = ""
        
    if "conversation_string" not in st.session_state:
        st.session_state["conversation_string"] = ""
        
    if "assessment" not in st.session_state:
        st.session_state["assessment"] = ""
        
    if "last_audio_size" not in st.session_state:
        st.session_state["last_audio_size"] = 0
        
    if "h_and_p" not in st.session_state:
        st.session_state["h_and_p"] = ""

    input_source = st.sidebar.radio("Choose to type or speak!", ("Text", "Microphone"), index=0)
    st.session_state.audio_off = st.sidebar.checkbox("Turn off voice response", value=False) 
    conversation_str = extracted_section + "**Learner Tasks:**\n\n" + st.session_state.learner_tasks + "\n\n" + "______" + "\n\n" + "**Clinical Interview:**\n\n"
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message(message["role"], avatar="👩‍⚕️"):
                st.markdown(message["content"])
                conversation_str += "👩‍⚕️: " + message["content"] + "\n\n"
        elif message["role"] == "assistant":
            with st.chat_message(message["role"], avatar="🤒"):
                st.markdown(message["content"])
                conversation_str += "🤒: " + message["content"] + "\n\n"
    conversation_str += "______" + "\n\n" + "**Orders:**\n\n" + st.session_state.orders_placed +  "**Results:**\n\n" + st.session_state.results + "\n\n"
    st.session_state.conversation_string = conversation_str

    if input_source == "Text":
        if prompt := st.chat_input("What's up?"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user", avatar="👩‍⚕️"):
                st.markdown(prompt)
            with st.chat_message("assistant", avatar="🤒"):    
                if st.session_state.model == "llama3-70b-8192":    
                    stream = groq_client.chat.completions.create(
                        model=st.session_state["model"],
                        messages=[
                            {"role": m["role"], "content": m["content"]}
                            for m in st.session_state.messages
                        ],
                        temperature=0.3,
                        stream=True,
                    )
                    st.session_state.sim_response = st.write_stream(parse_groq_stream(stream))
                elif st.session_state.model == "gpt-4o":
                    api_key = st.secrets["OPENAI_API_KEY"]
                    client = OpenAI(base_url="https://api.openai.com/v1", api_key=api_key)
                    completion = client.chat.completions.create(
                        model = st.session_state.model,
                        messages = [
                            {"role": m["role"], "content": m["content"]}
                            for m in st.session_state.messages
                        ],
                        temperature = 0.5,
                        max_tokens = 1000,
                        stream = True,   
                    )     
                    st.session_state.sim_response = st.write_stream(completion)
            st.session_state.messages.append({"role": "assistant", "content": st.session_state.sim_response})

            # Match the question with the checklist field and update it
            update_checklist_with_answer(prompt, st.session_state.sim_response, st.session_state.checklist_fields)

            # Update the checklist template with the new answers
            st.session_state.checklist_template = generate_checklist_template(st.session_state.checklist_fields)

    else:
        with st.sidebar:
            st.info("Click the green person-icon, pause 3 seconds, and begin to speak with natural speech.\
                    As soon as you pause, the LLM will start its response.")
            audio_bytes = audio_recorder(
                text="Click, pause, speak:",
                recording_color="#e8b62c",
                neutral_color="#6aa36f",
                icon_name="user",
                icon_size="3x",
            )

        if audio_bytes:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
                    temp_file.write(audio_bytes)
                    audio_file_path = temp_file.name
                
                file_stats = os.stat(audio_file_path)
                if st.session_state["last_audio_size"] != file_stats.st_size:
                    with st.spinner("Transcribing audio... Please wait."):
                        prompt = transcribe_audio(audio_file_path)
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    with st.chat_message("user", avatar="👩‍⚕️"):
                        st.markdown(prompt)
            finally:
                if 'audio_file_path' in locals():
                    os.remove(audio_file_path)
            audio_bytes = None

            if st.session_state["last_audio_size"] != file_stats.st_size:
                with st.chat_message("assistant", avatar="🤒"):
                    with st.spinner("Answering... Please wait."):     
                        if st.session_state.model == "llama3-70b-8192":   
                            stream = groq_client.chat.completions.create(
                                model=st.session_state["model"],
                                messages=[
                                    {"role": m["role"], "content": m["content"]}
                                    for m in st.session_state.messages
                                ],
                                temperature=0.3,
                                stream=True,
                            )
                            st.session_state.sim_response = st.write_stream(parse_groq_stream(stream))
                        elif st.session_state.model == "gpt-4o":
                            api_key = st.secrets["OPENAI_API_KEY"]
                            client = OpenAI(base_url="https://api.openai.com/v1", api_key=api_key)
                            completion = client.chat.completions.create(
                                model = st.session_state.model,
                                messages = [
                                    {"role": m["role"], "content": m["content"]}
                                    for m in st.session_state.messages
                                ],
                                temperature = 0.5,
                                max_tokens = 1000,
                                stream = True,   
                            )     
                            st.session_state.sim_response = st.write_stream(completion)
                st.session_state.messages.append({"role": "assistant", "content": st.session_state.sim_response})

                # Match the question with the checklist field and update it
                update_checklist_with_answer(prompt, st.session_state.sim_response, st.session_state.checklist_fields)

                # Update the checklist template with the new answers
                st.session_state.checklist_template = generate_checklist_template(st.session_state.checklist_fields)
    
    if st.session_state.audio_off == False:
        if st.session_state.sim_response:
            with st.spinner("Synthesizing audio... Please wait."):
                talk_stream("tts-1", st.session_state.voice, st.session_state.sim_response)
            autoplay_local_audio("last_interviewer.mp3")
            st.info("Note - this is an AI synthesized voice.")            
            st.session_state.sim_response = "" 
            os.remove("last_interviewer.mp3") 

    st.sidebar.divider()
    st.sidebar.subheader("Chart Access")
    
    orders = st.sidebar.checkbox("Place Orders/Take Actions", value=False)
    if orders:
        with st.sidebar:
            # Ensure the 'suggestions_text' key exists in the session state
            if "suggestions_text" not in st.session_state:
                # Generating suggestions for the placeholder
                placeholder_prompt = f"Based on the following case details, list in one word the various medical tests to place, medications to take and lab tests to take:\n\n{st.session_state.final_case}"
                placeholder_messages = [{"role": "user", "content": placeholder_prompt}]
                
                with st.spinner("Generating suggestions..."):
                    placeholder_results = llm_call("openai/gpt-4o", placeholder_messages)
                
                # Get the suggestions to display in the expander
                if placeholder_results:
                    st.session_state.suggestions_text = placeholder_results['choices'][0]['message']['content']
                else:
                    st.session_state.suggestions_text = "No suggestions available at the moment."

            # Text input for order details without suggestions as placeholder
            order_details = st.text_input("Order Details", value="", placeholder="E.g., examine lungs, CXR, CBC, furosemide 40 mg IV x 1, consult cardiology, etc.", key="order")

            # Expander to show the suggestions
            with st.expander("Suggestions for Orders/Actions", expanded=False):
                suggestions = st.session_state.suggestions_text.split('\n\n')
                medical_tests = []
                medications = []
                lab_tests = []

                # Classify suggestions into categories
                for suggestion in suggestions:
                    if "medical" in suggestion.lower():
                        medical_tests.append(suggestion.strip())
                    elif "medication" in suggestion.lower() or "med" in suggestion.lower():
                        medications.append(suggestion.strip())
                    elif "lab" in suggestion.lower():
                        lab_tests.append(suggestion.strip())

                checked_suggestions = {
                    "Medical Tests": [],
                    "Medications": [],
                    "Lab Tests": []
                }

                st.subheader("Medical Tests")
                for i, suggestion in enumerate(medical_tests):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"medical_test_{i}"):
                            checked_suggestions["Medical Tests"].append(suggestion)
                
                st.subheader("Medications")
                for i, suggestion in enumerate(medications):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"medication_{i}"):
                            checked_suggestions["Medications"].append(suggestion)
                
                st.subheader("Lab Tests")
                for i, suggestion in enumerate(lab_tests):
                    if suggestion:
                        if st.checkbox(suggestion, key=f"lab_test_{i}"):
                            checked_suggestions["Lab Tests"].append(suggestion)

            if st.button("Submit Orders/Take Actions"):
                # Get the current date and time
                current_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                order_details_with_datetime = f"{order_details}\n\nDate and Time of Request: {current_datetime}"
                
                # Update session state with the new orders
                combined_suggestions = ""
                for category, suggestions_list in checked_suggestions.items():
                    if suggestions_list:  # Only add category if there are checked suggestions
                        combined_suggestions += "\n".join(suggestions_list) + "\n\n"
            
                st.session_state.orders_placed = order_details_with_datetime + "\n\n" + combined_suggestions

                # Prompt for orders/actions based on the case details and current orders
                prompt = orders_prompt.format(order_details=order_details, case_details=st.session_state.final_case, order_datetime=current_datetime, prior_results=st.session_state.results)
                orders_messages = [{"role": "user", "content": prompt}]
                with st.spinner("Transmitting Orders... Please wait."):
                    orders_results = llm_call("openai/gpt-4o", orders_messages)
                    
                st.session_state.results = orders_results['choices'][0]['message']['content'] + "\n\n" + st.session_state.results
                
            with st.expander("Completed Orders/Actions", expanded=False):
                st.write(st.session_state.orders_placed)
                
            with st.expander("All Results of Orders/Actions", expanded=False):
                st.write(st.session_state.results) 

    html2 = markdown2.markdown(st.session_state.conversation_string, extras=["tables"])
    with st.sidebar:
        h_and_p = st.checkbox("Generate a History and Physical (no A/P section)", value=False)
        if h_and_p:
            prompt = h_and_p_prompt.format(conversation_transcript=st.session_state.conversation_string)
            h_and_p_messages = [{"role": "user", "content": prompt}]
            if st.sidebar.button("Create the History and Physical"):
                with st.sidebar:
                    with st.spinner("Writing History and Physical... Please wait."):
                        try:
                            h_and_p_response = llm_call("anthropic/claude-3-sonnet", h_and_p_messages)
                        except Exception as e:
                            st.error("Error formulating history and physical. Here are the error details: " + str(e))
                st.session_state.h_and_p = h_and_p_response['choices'][0]['message']['content']
            if st.session_state.h_and_p:
                with st.expander("History and Physical", expanded=False):
                    st.write(st.session_state.h_and_p)
                html = markdown2.markdown(st.session_state.h_and_p, extras=["tables"])
                with st.sidebar:
                    if st.button("Generate H&P PDF file"):
                        transcript_to_pdf(html, 'h_and_p.pdf')
                        with open("h_and_p.pdf", "rb") as f:
                            st.download_button("Download H&P PDF", f, "h_and_p.pdf")
        
        st.divider()     
        if st.button("Generate Transcript PDF file"):
            transcript_to_pdf(html2, 'transcript.pdf')
            with open("transcript.pdf", "rb") as f:
                st.download_button("Download Transcript PDF", f, "transcript.pdf")    
        assess = st.checkbox("Assess Interaction", value=False)
    
    if assess:
        student_level = st.sidebar.selectbox("Student Level", ["1st Year Medical Student", "2nd Year Medical Student", "3rd Year Medical Student", "4th Year Medical Student"])
        prompt = assessment_prompt.format(learner_tasks=st.session_state.learner_tasks, student_level=student_level, case_details=st.session_state.final_case, conversation_transcript=st.session_state.conversation_string, orders_placed=st.session_state.orders_placed, results=st.session_state.results)
        assessment_messages = [{"role": "user", "content": prompt}]
        if st.sidebar.button("Formulate Assessment"):
            with st.sidebar:
                with st.spinner("Formulating Assessment... Please wait."):
                    try:
                        assessment_response = llm_call("anthropic/claude-3-sonnet", assessment_messages)
                    except Exception as e:
                        st.error("Error formulating assessment, be sure to download the transcript and try again. Here are the error details: " + str(e))
            st.session_state.assessment = assessment_response['choices'][0]['message']['content']
        if st.session_state.assessment:
            with st.expander("Assessment", expanded=False):
                st.write(st.session_state.assessment)
            html = markdown2.markdown(st.session_state.assessment, extras=["tables"])
            with st.sidebar:
                if st.button("Generate Assessment PDF file"):
                    transcript_to_pdf(html, 'assessment.pdf')
                    with open("assessment.pdf", "rb") as f:
                        st.download_button("Download Assessment PDF", f, "assessment.pdf")

        st.divider()
        if st.session_state.checklist_template:
            checklist_html = markdown2.markdown(st.session_state.checklist_template, extras=["tables"])
            pdf_path = html_to_pdf(checklist_html, 'updated_checklist.pdf')
            if pdf_path:
                with open(pdf_path, "rb") as f:
                    st.download_button("Download Updated Checklist PDF", f, "updated_checklist.pdf")
#sprint3
with st.sidebar:
            selected_option = st.selectbox(
                "Select an option",
                ["Combined Word Doc", "Combined PDF"]
            )

            if st.button("Generate Combined Doc"):
                checklist_html = markdown2.markdown(st.session_state.get("checklist_template", ""), extras=["tables"])
                assessment_html = markdown2.markdown(st.session_state.get("assessment", ""), extras=["tables"]) if st.session_state.get("assessment") else ""
                orders_html = markdown2.markdown(st.session_state.get("orders_placed", ""), extras=["tables"]) if st.session_state.get("orders_placed") else ""

                if selected_option == "Combined Word Doc":
                    combined_doc_path = generate_combined_doc(checklist_html, assessment_html, orders_html, 'combined.docx')
                    if combined_doc_path:
                        with open(combined_doc_path, "rb") as f:
                            st.download_button("Download Combined Document", f, "combined.docx")

                elif selected_option == "Combined PDF":
                    combined_pdf_path = generate_combined_pdf(checklist_html, assessment_html, orders_html, 'combined.pdf')
                    if combined_pdf_path:
                        with open(combined_pdf_path, "rb") as f:
                            st.download_button("Download Combined Document", f, "combined.pdf")